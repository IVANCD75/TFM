"""
Generador de informe pericial siguiendo la estructura UNE 197010:2015.

"""
import io
import os
from datetime import datetime

from utils import resolve_path, get_logger, format_dt, format_size, safe_unix_dt
import report_content as RC

log = get_logger()


# --- HELPERS COMUNES DE FORMATO ---
# Los formateadores viven en utils.py (única fuente de verdad). Se mantienen
# estos alias locales para no reescribir las múltiples referencias del módulo.
_fmt_dt = format_dt
_fmt_size = format_size


def _bool_text(b):
    if b is True:
        return "Sí"
    if b is False:
        return "No"
    return "Sin datos"


# --- CONTENIDO COMÚN A AMBOS FORMATOS ---
def _construir_secciones(resultados, metadatos):
    """Construye una estructura de secciones reutilizable por PDF/DOCX."""
    perito = metadatos.get("perito", {})
    caso = metadatos.get("caso", {})

    # --- Resumen ejecutivo ---
    hashes = resultados.get("hashes") or {}
    n_usb = len(resultados.get("usb", []))
    actividad = resultados.get("actividad", {})
    n_prefetch = len(actividad.get("prefetch", []))
    seguridad = resultados.get("seguridad", {})
    n_logon_fail = len(seguridad.get("eventos", {}).get("logon_failed", []))
    n_audit_cleared = len(seguridad.get("eventos", {}).get("audit_cleared", []))
    apps = resultados.get("apps", {})
    n_apps = len(apps.get("instaladas", []))
    n_af = len(apps.get("anti_forensic", []))
    navegacion = resultados.get("navegacion", {})
    n_urls = len(navegacion.get("urls", []))
    n_kw = len(navegacion.get("keywords", []))

    # --- Hallazgos relevantes (texto) ---
    hallazgos = []
    if n_af:
        hallazgos.append(
            f"Se han detectado {n_af} aplicaciones potencialmente anti-forenses instaladas "
            f"en el equipo. Esto es un indicador relevante en escenarios de fuga de información."
        )
    if n_audit_cleared:
        hallazgos.append(
            f"Se han identificado {n_audit_cleared} eventos de borrado del registro de auditoría "
            f"(Event ID 1102), lo cual constituye un indicador clásico de manipulación intencional."
        )
    if n_logon_fail > 50:
        hallazgos.append(
            f"Se han registrado {n_logon_fail} intentos de inicio de sesión fallidos, "
            f"compatibles con un posible ataque de fuerza bruta."
        )
    if n_usb:
        hallazgos.append(
            f"Se identifican {n_usb} dispositivos USB conectados al equipo a lo largo de su uso."
        )
    if n_kw:
        hallazgos.append(
            f"Se han recuperado {n_kw} consultas en motores de búsqueda que pueden aportar contexto "
            f"sobre la intencionalidad del usuario."
        )
    if not hallazgos:
        hallazgos.append("No se identifican indicadores anómalos relevantes en el análisis preliminar.")

    return {
        "perito": perito,
        "caso": caso,
        "fecha_emision": metadatos.get("fecha_emision", datetime.now().strftime("%Y-%m-%d")),
        "ruta_imagen": resultados.get("ruta_imagen", ""),
        "hashes": hashes,
        "particiones": resultados.get("particiones", []),
        "sistema": {
            "hostname": resultados.get("hostname"),
            "os_name": resultados.get("os_name"),
            "release": resultados.get("release"),
            "build": resultados.get("build"),
            "owner": resultados.get("owner"),
            "install_date": safe_unix_dt(resultados.get("install_date_unix")),
            "timezone": resultados.get("timezone"),
            "last_shutdown": _fmt_dt(resultados.get("last_shutdown")),
        },
        "resumen_metricas": {
            "Dispositivos USB": n_usb,
            "Archivos Prefetch": n_prefetch,
            "Apps instaladas": n_apps,
            "Apps anti-forenses": n_af,
            "Logons fallidos": n_logon_fail,
            "URLs navegadas": n_urls,
            "Búsquedas web": n_kw,
        },
        "hallazgos": hallazgos,
        "usuarios": resultados.get("users", []),
        "apps_anti_forensic": apps.get("anti_forensic", []),
        "usb": resultados.get("usb", []),
        "actividad": actividad,
        "persistencia": resultados.get("persistencia", {}),
        "red": resultados.get("red", {}),
        "actividad_usuario": resultados.get("actividad_usuario", {}),
        "navegacion": navegacion,
        "seguridad": seguridad,
    }


# --- GENERADOR DE PDF ---
def generar_pdf(resultados, metadatos):
    """Genera el informe pericial en PDF. Devuelve bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        Table, TableStyle, KeepTogether, Image,
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.utils import ImageReader

    sec = _construir_secciones(resultados, metadatos)
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=3.2 * cm, bottomMargin=2.5 * cm,
        title=f"Informe Pericial — Caso {sec['caso'].get('referencia','S/R')}",
        author=sec["perito"].get("nombre", "Perito"),
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"],
                                 fontSize=22, leading=26, spaceAfter=20, alignment=TA_CENTER)
    h1 = ParagraphStyle("H1X", parent=styles["Heading1"],
                        fontSize=14, leading=18, spaceBefore=12, spaceAfter=8,
                        textColor=colors.HexColor("#1f4e79"))
    h2 = ParagraphStyle("H2X", parent=styles["Heading2"],
                        fontSize=12, leading=16, spaceBefore=8, spaceAfter=4,
                        textColor=colors.HexColor("#2e75b6"))
    body = ParagraphStyle("BodyX", parent=styles["BodyText"],
                          fontSize=10, leading=14, alignment=TA_JUSTIFY, spaceAfter=6)
    small = ParagraphStyle("SmallX", parent=body, fontSize=8, leading=10)
    centered = ParagraphStyle("CenterX", parent=body, alignment=TA_CENTER)

    table_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])

    def _table(headers, rows, col_widths=None):
        if not rows:
            return Paragraph("<i>Sin datos.</i>", small)
        data = [headers] + [[Paragraph(str(c) if c not in (None, "") else "—", small) for c in r] for r in rows]
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(table_style)
        return t

    story = []

    # --- PORTADA ---
    # Logo grande centrado (si existe)
    _logo_portada = resolve_path("img/Logo.png")
    if _logo_portada:
        try:
            _img_reader = ImageReader(_logo_portada)
            _iw, _ih = _img_reader.getSize()
            _ratio = (_ih / _iw) if _iw else 1.0
            _logo_w = 3.5 * cm
            _logo_h = _logo_w * _ratio
            story.append(Spacer(1, 1.5 * cm))
            _logo_img = Image(_logo_portada, width=_logo_w, height=_logo_h)
            _logo_img.hAlign = "CENTER"
            story.append(_logo_img)
            story.append(Spacer(1, 1 * cm))
        except Exception:
            story.append(Spacer(1, 4 * cm))
    else:
        story.append(Spacer(1, 4 * cm))

    story.append(Paragraph("INFORME PERICIAL INFORMÁTICO", title_style))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(
        f"<b>{sec['caso'].get('titulo', 'Análisis forense de evidencia digital')}</b>",
        centered,
    ))
    story.append(Spacer(1, 2 * cm))

    portada_data = [
        ["Referencia del informe:", sec["caso"].get("referencia", "—")],
        ["Procedimiento:", sec["caso"].get("procedimiento", "—")],
        ["Solicitante:", sec["caso"].get("solicitante", "—")],
        ["Juzgado / Órgano:", sec["caso"].get("juzgado", "—")],
        ["Perito:", sec["perito"].get("nombre", "—")],
        ["Titulación:", sec["perito"].get("titulacion", "—")],
        ["Nº Colegiado:", sec["perito"].get("colegiado", "—")],
        ["Fecha de emisión:", sec["fecha_emision"]],
    ]
    t = Table(portada_data, colWidths=[5 * cm, 11 * cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.lightgrey),
    ]))
    story.append(t)
    story.append(PageBreak())

    # --- DECLARACIÓN DE IMPARCIALIDAD ---
    story.append(Paragraph("1. Declaración de imparcialidad y juramento", h1))
    nombre = sec["perito"].get("nombre", "[Nombre del perito]")
    tit = sec["perito"].get("titulacion", "[Titulación]")
    col = sec["perito"].get("colegiado", "[Nº colegiado]")
    story.append(Paragraph(
        RC.DECLARACION_INTRO.format(
            nombre=f"<b>{nombre}</b>", titulacion=f"<b>{tit}</b>",
            colegiado=f"<b>{col}</b>",
        ),
        body,
    ))
    for d in RC.DECLARACIONES_IMPARCIALIDAD:
        story.append(Paragraph(f"• {d}", body))

    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(
        f"En <b>{sec['perito'].get('lugar_firma', '________________')}</b>, a "
        f"<b>{sec['fecha_emision']}</b>.",
        body,
    ))
    story.append(Spacer(1, 2 * cm))
    story.append(Paragraph("Fdo.: ____________________________", body))
    story.append(Paragraph(f"{nombre}", body))
    story.append(PageBreak())

    # --- ÍNDICE ---
    story.append(Paragraph("Índice", h1))
    for item in RC.INDICE_ITEMS:
        story.append(Paragraph(item, body))
    story.append(PageBreak())

    # --- RESUMEN EJECUTIVO ---
    story.append(Paragraph("2. Resumen ejecutivo", h1))
    story.append(Paragraph(RC.RESUMEN_EJECUTIVO, body))
    story.append(Paragraph("<b>Métricas principales:</b>", body))
    metricas_rows = [[k, str(v)] for k, v in sec["resumen_metricas"].items()]
    story.append(_table(["Categoría", "Valor"], metricas_rows, [10 * cm, 6 * cm]))

    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("<b>Hallazgos relevantes:</b>", body))
    for h in sec["hallazgos"]:
        story.append(Paragraph(f"• {h}", body))
    story.append(PageBreak())

    # --- OBJETO DEL ENCARGO ---
    story.append(Paragraph("3. Objeto del encargo", h1))
    objeto_descrito = sec["caso"].get("objeto") or RC.OBJETO_DEFECTO
    story.append(Paragraph(objeto_descrito, body))
    story.append(PageBreak())

    # --- ANTECEDENTES ---
    story.append(Paragraph("4. Antecedentes", h1))
    antecedentes = sec["caso"].get("antecedentes") or RC.ANTECEDENTES_DEFECTO
    story.append(Paragraph(antecedentes, body))
    story.append(PageBreak())

    # --- FUENTES DE INFORMACIÓN ---
    story.append(Paragraph("5. Fuentes de información", h1))
    story.append(Paragraph(RC.FUENTES_INTRO, body))

    # Ya que la ruta del archivo puede ser muy larga y he comprobado que se sale de la tabla al generar el pdf:
    styles = getSampleStyleSheet()
    ruta_archivo = Paragraph(
        sec["ruta_imagen"] or "—",
        styles["BodyText"]
    )

    fuentes_rows = [
        ["Ruta del archivo", ruta_archivo],
        ["Tamaño de la imagen", _fmt_size(sec["hashes"].get("tamaño_bytes", 0)) if sec["hashes"] else "—"],
        ["Origen del hash", sec["hashes"].get("fuente", "—") if sec["hashes"] else "—"],
        ["MD5", sec["hashes"].get("MD5") or "No calculado" if sec["hashes"] else "—"],
        ["SHA-1", sec["hashes"].get("SHA-1") or "No calculado" if sec["hashes"] else "—"],
        ["SHA-256", sec["hashes"].get("SHA-256") or "No calculado" if sec["hashes"] else "—"],
    ]
    t = Table(fuentes_rows, colWidths=[4 * cm, 12 * cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#f7f7f7"), colors.white]),
    ]))
    story.append(t)

    if sec["particiones"]:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("<b>Estructura de particiones:</b>", body))
        rows = [
            [p["Nº"], p["Descripción"][:40], p["Sistema Archivos"],
             f"{p['Sector Inicio']:,}", f"{p['Total Sectores']:,}", p["Tamaño"]]
            for p in sec["particiones"]
        ]
        story.append(_table(
            ["Nº", "Descripción", "FS", "Sector Inicio", "Total Sectores", "Tamaño"],
            rows,
        ))
    story.append(PageBreak())

    # --- ANÁLISIS ---
    story.append(Paragraph("6. Análisis", h1))
    story.append(Paragraph(RC.ANALISIS_INTRO, body))

    # --- SISTEMA ---
    story.append(Paragraph("6.1 Identificación del sistema", h2))
    sys_rows = [[k, v if v else "—"] for k, v in [
        ("Hostname", sec["sistema"]["hostname"]),
        ("Sistema operativo", sec["sistema"]["os_name"]),
        ("Versión / Build", f"{sec['sistema']['release']} / {sec['sistema']['build']}"),
        ("Propietario registrado", sec["sistema"]["owner"]),
        ("Fecha de instalación", sec["sistema"]["install_date"]),
        ("Zona horaria", sec["sistema"]["timezone"]),
        ("Último apagado registrado", sec["sistema"]["last_shutdown"]),
    ]]
    story.append(_table(["Propiedad", "Valor"], sys_rows, [6 * cm, 10 * cm]))

    # --- USUARIOS ---
    story.append(Paragraph("6.2 Cuentas de usuario", h2))
    if sec["usuarios"]:
        story.append(_table(["Nombre"], [[u] for u in sec["usuarios"]]))
    else:
        story.append(Paragraph("<i>No se pudo extraer información de la hive SAM.</i>", body))

    # --- APPS ---
    story.append(Paragraph("6.3 Aplicaciones instaladas", h2))
    apps = resultados.get("apps", {})
    n_apps = len(apps.get("instaladas", []))
    n_af = len(apps.get("anti_forensic", []))
    story.append(Paragraph(
        f"Se han identificado <b>{n_apps}</b> aplicaciones instaladas en el sistema, "
        f"de las cuales <b>{n_af}</b> se corresponden con herramientas potencialmente "
        f"anti-forenses.",
        body,
    ))
    if sec["apps_anti_forensic"]:
        story.append(Paragraph("<b>Aplicaciones de interés (anti-forenses):</b>", body))
        rows = [
            [a["Nombre"], a.get("Versión", ""), a.get("Editor", ""),
             a.get("Fecha Instalación", ""), a.get("Indicador", "")]
            for a in sec["apps_anti_forensic"]
        ]
        story.append(_table(
            ["Nombre", "Versión", "Editor", "Fecha Inst.", "Indicador"],
            rows,
        ))
    if apps.get("instaladas"):
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph("<b>Aplicaciones instaladas (lista completa):</b>", body))
        rows = [
            [a["Nombre"][:40], a.get("Versión", "")[:18], a.get("Editor", "")[:30],
             a.get("Fecha Instalación", "")]
            for a in apps["instaladas"][:60]
        ]
        story.append(_table(
            ["Nombre", "Versión", "Editor", "Fecha Inst."],
            rows,
        ))
        if len(apps["instaladas"]) > 60:
            story.append(Paragraph(
                f"<i>Se muestran las primeras 60 entradas de {len(apps['instaladas'])} totales.</i>",
                small,
            ))

    # --- USB ---
    story.append(PageBreak())
    story.append(Paragraph("6.4 Historial de dispositivos USB", h2))
    if sec["usb"]:
        story.append(Paragraph(
            f"Se identifican <b>{len(sec['usb'])}</b> dispositivos USB en el historial "
            f"del equipo (clave USBSTOR del registro SYSTEM cruzada con setupapi.dev.log).",
            body,
        ))
        rows = [
            [u["Fabricante"][:20], u["Modelo"][:25], u["Nº Serie"][:25],
             _fmt_dt(u["Primera Conexión"]), _fmt_dt(u["Última Conexión"])]
            for u in sec["usb"]
        ]
        story.append(_table(
            ["Fabricante", "Modelo", "Nº Serie", "Primera Conexión", "Última Conexión"],
            rows,
        ))
    else:
        story.append(Paragraph("<i>No se han registrado dispositivos USB.</i>", body))

    # --- ACTIVIDAD DE EJECUCIÓN ---
    story.append(Paragraph("6.5 Actividad de ejecución", h2))
    actividad = sec["actividad"]
    if actividad.get("prefetch") or actividad.get("userassist"):
        if actividad.get("prefetch"):
            story.append(Paragraph(
                f"<b>Prefetch:</b> {len(actividad['prefetch'])} archivos analizados. "
                f"Top 15 ejecuciones más recientes:",
                body,
            ))
            rows = [
                [p["Ejecutable"][:50], _fmt_dt(p["Última Ejecución"])]
                for p in actividad["prefetch"][:15]
            ]
            story.append(_table(["Ejecutable", "Última Ejecución"], rows, [10 * cm, 6 * cm]))

        if actividad.get("userassist"):
            story.append(Paragraph(
                f"<b>UserAssist:</b> {len(actividad['userassist'])} entradas. "
                f"Top 15 aplicaciones más utilizadas:",
                body,
            ))
            top = sorted(actividad["userassist"], key=lambda x: x["Veces Ejecutado"], reverse=True)[:15]
            rows = [
                [u["Usuario"], u["Aplicación"][:40], u["Veces Ejecutado"], _fmt_dt(u["Última Ejecución"])]
                for u in top
            ]
            story.append(_table(
                ["Usuario", "Aplicación", "Veces", "Última Ejec."],
                rows,
            ))
    else:
        story.append(Paragraph("<i>No se obtuvieron datos de actividad.</i>", body))

    # --- PERSISTENCIA ---
    story.append(PageBreak())
    story.append(Paragraph("6.6 Persistencia", h2))
    persistencia = sec["persistencia"]
    runs = persistencia.get("run_keys", [])
    tasks = persistencia.get("scheduled_tasks", [])
    sospech_runs = [r for r in runs if r.get("Sospechoso")]
    sospech_tasks = [t for t in tasks if t.get("Sospechoso")]
    story.append(Paragraph(
        f"Se identifican <b>{len(runs)}</b> entradas Run/RunOnce y "
        f"<b>{len(tasks)}</b> tareas programadas. De estas, "
        f"<b>{len(sospech_runs)}</b> y <b>{len(sospech_tasks)}</b> respectivamente han sido "
        f"marcadas como potencialmente sospechosas por la heurística de la herramienta.",
        body,
    ))
    if sospech_runs:
        story.append(Paragraph("<b>Run keys sospechosas:</b>", body))
        rows = [[r["Origen"], r["Usuario"], r["Nombre"][:30], r["Comando"][:60]] for r in sospech_runs]
        story.append(_table(["Origen", "Usuario", "Nombre", "Comando"], rows))
    if sospech_tasks:
        story.append(Paragraph("<b>Tareas programadas sospechosas:</b>", body))
        rows = [[t["Nombre"][:30], t["Autor"][:20], t["Creada"], t["Acción"][:60]] for t in sospech_tasks]
        story.append(_table(["Nombre", "Autor", "Creada", "Acción"], rows))

    # --- RED ---
    story.append(Paragraph("6.7 Actividad de red", h2))
    red = sec["red"]
    if red.get("profiles"):
        story.append(Paragraph(f"<b>Perfiles de red ({len(red['profiles'])}):</b>", body))
        rows = [
            [p["Nombre (SSID)"][:30], p["Tipo"], _fmt_dt(p["Creada"]), _fmt_dt(p["Última Conexión"])]
            for p in red["profiles"]
        ]
        story.append(_table(["SSID/Nombre", "Tipo", "Creada", "Última Conexión"], rows))
    if red.get("interfaces"):
        story.append(Paragraph(f"<b>Interfaces IP activas ({len(red['interfaces'])}):</b>", body))
        rows = [
            [i["Modo"], i["IP"], i["Máscara"], i["Gateway"], i["DNS"][:40]]
            for i in red["interfaces"]
        ]
        story.append(_table(["Modo", "IP", "Máscara", "Gateway", "DNS"], rows))

    # --- ACTIVIDAD USUARIO ---
    story.append(PageBreak())
    story.append(Paragraph("6.8 Actividad del usuario", h2))
    au = sec["actividad_usuario"]
    rb = au.get("recycle_bin", {})
    story.append(Paragraph(
        f"<b>Papelera de reciclaje:</b> {rb.get('total_count', 0)} archivos "
        f"({_fmt_size(rb.get('total_size', 0))}).",
        body,
    ))
    if rb.get("files"):
        rows = [
            [f["Archivo Original"], _fmt_size(f["Tamaño (bytes)"]), _fmt_dt(f["Fecha Borrado"])]
            for f in rb["files"][:20]
        ]
        story.append(_table(["Archivo", "Tamaño", "Fecha Borrado"], rows))

    network_mru = au.get("network_mru", {})
    if network_mru.get("mapped"):
        story.append(Paragraph(
            f"<b>Drives de red mapeados ({len(network_mru['mapped'])}):</b>", body
        ))
        rows = [
            [m.get("Usuario", ""), m.get("Letra", ""), m.get("Ruta UNC", "")[:50],
             _fmt_dt(m.get("Última Modif."))]
            for m in network_mru["mapped"]
        ]
        story.append(_table(["Usuario", "Letra", "Ruta UNC", "Última Modif."], rows))

    # --- NAVEGACIÓN ---
    story.append(Paragraph("6.9 Navegación web y búsquedas", h2))
    nav = sec["navegacion"]
    story.append(Paragraph(
        f"<b>{len(nav.get('urls', []))} URLs</b> recuperadas del historial, "
        f"<b>{len(nav.get('keywords', []))} consultas en buscadores</b> y "
        f"<b>{len(nav.get('downloads', []))} descargas</b>.",
        body,
    ))
    if nav.get("keywords"):
        story.append(Paragraph("<b>Términos de búsqueda detectados (top 25):</b>", body))
        rows = [
            [k["Usuario"], k["Buscador"][:20], k["Término"][:60], _fmt_dt(k["Última Búsqueda"])]
            for k in nav["keywords"][:25]
        ]
        story.append(_table(["Usuario", "Buscador", "Término", "Última Búsqueda"], rows))
    if nav.get("downloads"):
        story.append(Paragraph("<b>Descargas (top 15):</b>", body))
        rows = [
            [d["Usuario"], (d.get("Archivo") or "")[-50:], _fmt_size(d["Tamaño (bytes)"]), _fmt_dt(d["Inicio"])]
            for d in nav["downloads"][:15]
        ]
        story.append(_table(["Usuario", "Archivo", "Tamaño", "Inicio"], rows))

    # --- SEGURIDAD ---
    story.append(PageBreak())
    story.append(Paragraph("6.10 Estado de seguridad", h2))
    seg = sec["seguridad"]
    defender = seg.get("defender", {})
    firewall = seg.get("firewall", {})
    eventos = seg.get("eventos", {})

    rows = [
        ["Defender — protección en tiempo real", _bool_text(defender.get("rt_monitoring"))],
        ["Defender — Tamper Protection", _bool_text(defender.get("tamper_protection"))],
        ["AV de terceros detectados", ", ".join(defender.get("third_party_av", [])) or "Ninguno"],
        ["Firewall — Perfil Dominio", _bool_text(firewall.get("domain"))],
        ["Firewall — Perfil Privado", _bool_text(firewall.get("private"))],
        ["Firewall — Perfil Público", _bool_text(firewall.get("public"))],
        ["Eventos analizados", str(eventos.get("total_revisados", 0))],
        ["Logons exitosos (4624)", str(len(eventos.get("logon_success", [])))],
        ["Logons fallidos (4625)", str(len(eventos.get("logon_failed", [])))],
        ["Borrado auditoría (1102)", str(len(eventos.get("audit_cleared", [])))],
    ]
    story.append(_table(["Categoría", "Estado"], rows, [10 * cm, 6 * cm]))

    if eventos.get("audit_cleared"):
        story.append(Paragraph(f"<b>{RC.EVENTO_1102_NOTA}</b>", body))

    # --- CONCLUSIONES ---
    story.append(PageBreak())
    story.append(Paragraph("7. Conclusiones", h1))
    story.append(Paragraph(RC.CONCLUSIONES_INTRO, body))
    for i, h in enumerate(sec["hallazgos"], 1):
        story.append(Paragraph(f"{i}. {h}", body))

    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(RC.CONCLUSIONES_TRIAJE, body))
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(
        RC.CIERRE_PERICIAL.format(
            lugar_firma=f"<b>{sec['perito'].get('lugar_firma', '_____________')}</b>",
            fecha_emision=f"<b>{sec['fecha_emision']}</b>",
        ),
        body,
    ))
    story.append(Spacer(1, 1.5 * cm))
    story.append(Paragraph(f"Fdo.: {nombre}", body))
    story.append(Paragraph(f"{tit} — Nº Colegiado: {col}", body))

    # --- ANEXOS ---
    story.append(PageBreak())
    story.append(Paragraph("8. Anexos", h1))
    story.append(Paragraph(RC.ANEXOS_INTRO, body))
    story.append(Paragraph("<b>Glosario rápido:</b>", body))
    for term, desc in RC.GLOSARIO:
        story.append(Paragraph(f"<b>{term}:</b> {desc}", body))

    # Header (logo arriba-izq) + footer (paginación abajo-dcha) en interiores.
    # La portada (página 1) se trata distinto: el logo va grande y centrado.
    logo_path = resolve_path("img/Logo.png")
    page_decorator = _make_page_decorator(logo_path)
    doc.build(story, onFirstPage=page_decorator, onLaterPages=page_decorator)
    buffer.seek(0)
    return buffer.getvalue()


def _make_page_decorator(logo_path):
    """Crea la función de decoración de página (header + footer) con el logo.

    - Páginas interiores (>=2): logo pequeño arriba-izquierda + línea de
      referencia centrada + número de página abajo-derecha.
    - Portada (página 1): no se dibuja el header (el logo va en el flujo
      del documento, grande y centrado), solo un footer discreto.
    """
    from reportlab.lib.utils import ImageReader

    logo_reader = None
    logo_ratio = 1.0
    if logo_path:
        try:
            logo_reader = ImageReader(logo_path)
            iw, ih = logo_reader.getSize()
            logo_ratio = (ih / iw) if iw else 1.0
        except Exception:
            logo_reader = None

    cm = 28.35  # 1 cm en puntos

    def _decorate(canvas, doc):
        canvas.saveState()
        page_width = doc.pagesize[0]
        page_height = doc.pagesize[1]

        # --- HEADER: logo arriba-izquierda (solo páginas interiores) ---
        if doc.page > 1 and logo_reader is not None:
            logo_w = 1.1 * cm
            logo_h = logo_w * logo_ratio
            x = 2.5 * cm
            # Colocar el logo dentro de la banda de margen superior, por
            # encima del área de contenido (que empieza a 2.5 cm del borde),
            # de modo que no se solape con el primer título de la página.
            y = page_height - 1.5 * cm
            try:
                canvas.drawImage(
                    logo_reader, x, y, width=logo_w, height=logo_h,
                    preserveAspectRatio=True, mask="auto",
                )
            except Exception:
                pass

        # --- FOOTER ---
        canvas.setFont("Helvetica", 8)
        canvas.setFillGray(0.4)
        canvas.drawString(
            2.5 * cm, 1.5 * cm,
            "Informe Pericial — Triaje Forense Automatizado",
        )
        canvas.drawRightString(
            page_width - 2.5 * cm, 1.5 * cm,
            f"Página {doc.page}",
        )
        canvas.restoreState()

    return _decorate



# --- GENERADOR DOCX ---
def _docx_set_page_number_field(paragraph):
    """Inserta un campo dinámico de número de página en un párrafo DOCX."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = paragraph.add_run()

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _docx_configurar_header_footer(section, logo_path):
    """Configura, para una sección Word:
       - Header de páginas interiores: logo arriba-izquierda.
       - Footer de páginas interiores: referencia (izq) + 'Página N' (dcha).
       - Portada (first page): header/footer vacíos.
    """
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

    # --- HEADER interior con logo ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path:
        try:
            from docx.shared import Inches
            run = hp.add_run()
            run.add_picture(logo_path, width=Inches(0.5))
        except Exception:
            pass

    # --- FOOTER interior con paginación ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    # Tabs: texto a la izquierda, número a la derecha
    usable_width = section.page_width - section.left_margin - section.right_margin
    fp.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    run_ref = fp.add_run("Informe Pericial — Triaje Forense Automatizado\t")
    run_ref.font.size = Pt(8)
    run_ref.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run_pag = fp.add_run("Página ")
    run_pag.font.size = Pt(8)
    run_pag.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _docx_set_page_number_field(fp)

    # --- PORTADA (first page): vacíos ---
    fph = section.first_page_header
    fph.is_linked_to_previous = False
    if fph.paragraphs:
        fph.paragraphs[0].text = ""
    fpf = section.first_page_footer
    fpf.is_linked_to_previous = False
    if fpf.paragraphs:
        fpf.paragraphs[0].text = ""


def generar_docx(resultados, metadatos):
    """Genera el informe pericial en formato Word. Devuelve bytes."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    sec = _construir_secciones(resultados, metadatos)
    doc = Document()

    # --- Estilos base ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    logo_path = resolve_path("img/Logo.png")

    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        # Primera página (portada) con header/footer propio y vacío
        section.different_first_page_header_footer = True
        _docx_configurar_header_footer(section, logo_path)

    def _add_h1(text):
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    def _add_h2(text):
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    def _add_p(text, bold=False, center=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        if bold:
            run.bold = True
        return p

    def _add_table(headers, rows):
        if not rows:
            doc.add_paragraph("Sin datos.", style="Intense Quote")
            return
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v) if v not in (None, "") else "—"
                for paragraph in cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # --- PORTADA ---
    # Logo grande centrado (si existe)
    if logo_path:
        try:
            from docx.shared import Inches
            logo_par = doc.add_paragraph()
            logo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = logo_par.add_run()
            run_logo.add_picture(logo_path, width=Inches(1.6))
            doc.add_paragraph()
        except Exception:
            for _ in range(4):
                doc.add_paragraph()
    else:
        for _ in range(4):
            doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("INFORME PERICIAL INFORMÁTICO")
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    _add_p(sec["caso"].get("titulo", "Análisis forense de evidencia digital"),
           bold=True, center=True)

    doc.add_paragraph()
    doc.add_paragraph()

    portada_data = [
        ("Referencia del informe", sec["caso"].get("referencia", "—")),
        ("Procedimiento", sec["caso"].get("procedimiento", "—")),
        ("Solicitante", sec["caso"].get("solicitante", "—")),
        ("Juzgado / Órgano", sec["caso"].get("juzgado", "—")),
        ("Perito", sec["perito"].get("nombre", "—")),
        ("Titulación", sec["perito"].get("titulacion", "—")),
        ("Nº Colegiado", sec["perito"].get("colegiado", "—")),
        ("Fecha de emisión", sec["fecha_emision"]),
    ]
    t = doc.add_table(rows=len(portada_data), cols=2)
    for i, (k, v) in enumerate(portada_data):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = str(v)
        for run in t.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_page_break()

    # --- DECLARACIÓN ---
    _add_h1("1. Declaración de imparcialidad y juramento")
    nombre = sec["perito"].get("nombre", "[Nombre]")
    tit = sec["perito"].get("titulacion", "[Titulación]")
    col = sec["perito"].get("colegiado", "[Nº colegiado]")
    _add_p(RC.DECLARACION_INTRO.format(nombre=nombre, titulacion=tit, colegiado=col))
    for d in RC.DECLARACIONES_IMPARCIALIDAD:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_paragraph()
    _add_p(f"En {sec['perito'].get('lugar_firma', '________')}, a {sec['fecha_emision']}.")
    doc.add_paragraph()
    doc.add_paragraph()
    _add_p("Fdo.: ____________________________")
    _add_p(nombre)
    doc.add_page_break()

    # --- RESUMEN EJECUTIVO ---
    _add_h1("2. Resumen ejecutivo")
    _add_p(RC.RESUMEN_EJECUTIVO)

    _add_p("Métricas principales:", bold=True)
    _add_table(["Categoría", "Valor"],
               [[k, str(v)] for k, v in sec["resumen_metricas"].items()])

    _add_p("Hallazgos relevantes:", bold=True)
    for h in sec["hallazgos"]:
        doc.add_paragraph(h, style="List Bullet")

    doc.add_page_break()

    # --- OBJETO ---
    _add_h1("3. Objeto del encargo")
    objeto_descrito = sec["caso"].get("objeto") or RC.OBJETO_DEFECTO
    _add_p(objeto_descrito)
    doc.add_page_break()

    # --- ANTECEDENTES ---
    _add_h1("4. Antecedentes")
    _add_p(sec["caso"].get("antecedentes") or RC.ANTECEDENTES_DEFECTO)
    doc.add_page_break()

    # --- FUENTES ---
    _add_h1("5. Fuentes de información")
    _add_p(RC.FUENTES_INTRO)
    h = sec["hashes"] or {}
    _add_table(["Propiedad", "Valor"], [
        ["Ruta del archivo", sec["ruta_imagen"] or "—"],
        ["Tamaño", _fmt_size(h.get("tamaño_bytes", 0))],
        ["Origen del hash", h.get("fuente", "—")],
        ["MD5", h.get("MD5") or "No calculado"],
        ["SHA-1", h.get("SHA-1") or "No calculado"],
        ["SHA-256", h.get("SHA-256") or "No calculado"],
    ])

    if sec["particiones"]:
        _add_p("Estructura de particiones:", bold=True)
        _add_table(
            ["Nº", "Descripción", "FS", "Sector Inicio", "Total Sectores", "Tamaño"],
            [[p["Nº"], p["Descripción"][:40], p["Sistema Archivos"],
              f"{p['Sector Inicio']:,}", f"{p['Total Sectores']:,}", p["Tamaño"]]
             for p in sec["particiones"]],
        )
    doc.add_page_break()

    # --- ANÁLISIS ---
    _add_h1("6. Análisis")
    _add_p(RC.ANALISIS_INTRO)

    _add_h2("6.1 Identificación del sistema")
    _add_table(["Propiedad", "Valor"], [
        ["Hostname", sec["sistema"]["hostname"]],
        ["Sistema operativo", sec["sistema"]["os_name"]],
        ["Versión / Build", f"{sec['sistema']['release']} / {sec['sistema']['build']}"],
        ["Propietario registrado", sec["sistema"]["owner"]],
        ["Fecha de instalación", sec["sistema"]["install_date"]],
        ["Zona horaria", sec["sistema"]["timezone"]],
        ["Último apagado registrado", sec["sistema"]["last_shutdown"]],
    ])

    _add_h2("6.2 Cuentas de usuario")
    if sec["usuarios"]:
        _add_table(["Nombre"], [[u] for u in sec["usuarios"]])

    _add_h2("6.3 Aplicaciones instaladas")
    apps = resultados.get("apps", {})
    _add_p(f"Total: {len(apps.get('instaladas', []))} aplicaciones. "
           f"De estas, {len(apps.get('anti_forensic', []))} son anti-forenses.")
    if sec["apps_anti_forensic"]:
        _add_p("Aplicaciones anti-forenses detectadas:", bold=True)
        _add_table(
            ["Nombre", "Versión", "Editor", "Fecha Inst.", "Indicador"],
            [[a["Nombre"], a.get("Versión", ""), a.get("Editor", ""),
              a.get("Fecha Instalación", ""), a.get("Indicador", "")]
             for a in sec["apps_anti_forensic"]],
        )
    if apps.get("instaladas"):
        _add_p("Lista completa (primeras 60):", bold=True)
        _add_table(
            ["Nombre", "Versión", "Editor", "Fecha Inst."],
            [[a["Nombre"][:40], a.get("Versión", "")[:18],
              a.get("Editor", "")[:30], a.get("Fecha Instalación", "")]
             for a in apps["instaladas"][:60]],
        )

    _add_h2("6.4 Historial de dispositivos USB")
    if sec["usb"]:
        _add_p(f"{len(sec['usb'])} dispositivos USB en el historial:")
        _add_table(
            ["Fabricante", "Modelo", "Nº Serie", "Primera Conexión", "Última Conexión"],
            [[u["Fabricante"][:20], u["Modelo"][:25], u["Nº Serie"][:25],
              _fmt_dt(u["Primera Conexión"]), _fmt_dt(u["Última Conexión"])]
             for u in sec["usb"]],
        )

    _add_h2("6.5 Actividad de ejecución")
    actividad = sec["actividad"]
    if actividad.get("prefetch"):
        _add_p(f"Prefetch: {len(actividad['prefetch'])} entradas. Top 15:", bold=True)
        _add_table(["Ejecutable", "Última Ejecución"],
                   [[p["Ejecutable"][:50], _fmt_dt(p["Última Ejecución"])]
                    for p in actividad["prefetch"][:15]])
    if actividad.get("userassist"):
        _add_p(f"UserAssist: {len(actividad['userassist'])} entradas. Top 15 más usadas:",
               bold=True)
        top = sorted(actividad["userassist"], key=lambda x: x["Veces Ejecutado"], reverse=True)[:15]
        _add_table(["Usuario", "Aplicación", "Veces", "Última Ejec."],
                   [[u["Usuario"], u["Aplicación"][:40], u["Veces Ejecutado"],
                     _fmt_dt(u["Última Ejecución"])] for u in top])

    _add_h2("6.6 Persistencia")
    p = sec["persistencia"]
    runs = p.get("run_keys", [])
    tasks = p.get("scheduled_tasks", [])
    sospech_runs = [r for r in runs if r.get("Sospechoso")]
    sospech_tasks = [t for t in tasks if t.get("Sospechoso")]
    _add_p(f"{len(runs)} entradas Run/RunOnce, {len(tasks)} tareas programadas. "
           f"Sospechosas: {len(sospech_runs)} runs, {len(sospech_tasks)} tareas.")
    if sospech_runs:
        _add_table(["Origen", "Usuario", "Nombre", "Comando"],
                   [[r["Origen"], r["Usuario"], r["Nombre"][:30], r["Comando"][:60]]
                    for r in sospech_runs])
    if sospech_tasks:
        _add_table(["Nombre", "Autor", "Creada", "Acción"],
                   [[t["Nombre"][:30], t["Autor"][:20], t["Creada"], t["Acción"][:60]]
                    for t in sospech_tasks])

    _add_h2("6.7 Actividad de red")
    red = sec["red"]
    if red.get("profiles"):
        _add_p(f"Perfiles de red ({len(red['profiles'])}):", bold=True)
        _add_table(["SSID/Nombre", "Tipo", "Creada", "Última Conexión"],
                   [[p["Nombre (SSID)"][:30], p["Tipo"], _fmt_dt(p["Creada"]),
                     _fmt_dt(p["Última Conexión"])] for p in red["profiles"]])
    if red.get("interfaces"):
        _add_p(f"Interfaces IP ({len(red['interfaces'])}):", bold=True)
        _add_table(["Modo", "IP", "Máscara", "Gateway", "DNS"],
                   [[i["Modo"], i["IP"], i["Máscara"], i["Gateway"], i["DNS"][:40]]
                    for i in red["interfaces"]])

    _add_h2("6.8 Actividad del usuario")
    au = sec["actividad_usuario"]
    rb = au.get("recycle_bin", {})
    _add_p(f"Papelera: {rb.get('total_count', 0)} archivos ({_fmt_size(rb.get('total_size', 0))}).")
    if rb.get("files"):
        _add_table(["Archivo", "Tamaño", "Fecha Borrado"],
                   [[f["Archivo Original"], _fmt_size(f["Tamaño (bytes)"]),
                     _fmt_dt(f["Fecha Borrado"])] for f in rb["files"][:20]])
    network_mru = au.get("network_mru", {})
    if network_mru.get("mapped"):
        _add_p(f"Drives de red mapeados ({len(network_mru['mapped'])}):", bold=True)
        _add_table(["Usuario", "Letra", "Ruta UNC", "Última Modif."],
                   [[m.get("Usuario", ""), m.get("Letra", ""), m.get("Ruta UNC", "")[:50],
                     _fmt_dt(m.get("Última Modif."))] for m in network_mru["mapped"]])

    _add_h2("6.9 Navegación web y búsquedas")
    nav = sec["navegacion"]
    _add_p(f"{len(nav.get('urls', []))} URLs, {len(nav.get('keywords', []))} consultas, "
           f"{len(nav.get('downloads', []))} descargas.")
    if nav.get("keywords"):
        _add_p("Búsquedas (top 25):", bold=True)
        _add_table(["Usuario", "Buscador", "Término", "Última Búsqueda"],
                   [[k["Usuario"], k["Buscador"][:20], k["Término"][:60],
                     _fmt_dt(k["Última Búsqueda"])] for k in nav["keywords"][:25]])
    if nav.get("downloads"):
        _add_p("Descargas (top 15):", bold=True)
        _add_table(["Usuario", "Archivo", "Tamaño", "Inicio"],
                   [[d["Usuario"], (d.get("Archivo") or "")[-50:],
                     _fmt_size(d["Tamaño (bytes)"]), _fmt_dt(d["Inicio"])]
                    for d in nav["downloads"][:15]])

    _add_h2("6.10 Estado de seguridad")
    seg = sec["seguridad"]
    defender = seg.get("defender", {})
    firewall = seg.get("firewall", {})
    eventos = seg.get("eventos", {})
    _add_table(["Categoría", "Estado"], [
        ["Defender — protección en tiempo real", _bool_text(defender.get("rt_monitoring"))],
        ["Defender — Tamper Protection", _bool_text(defender.get("tamper_protection"))],
        ["AV de terceros detectados", ", ".join(defender.get("third_party_av", [])) or "Ninguno"],
        ["Firewall — Dominio", _bool_text(firewall.get("domain"))],
        ["Firewall — Privado", _bool_text(firewall.get("private"))],
        ["Firewall — Público", _bool_text(firewall.get("public"))],
        ["Eventos analizados", str(eventos.get("total_revisados", 0))],
        ["Logons exitosos (4624)", str(len(eventos.get("logon_success", [])))],
        ["Logons fallidos (4625)", str(len(eventos.get("logon_failed", [])))],
        ["Borrado auditoría (1102)", str(len(eventos.get("audit_cleared", [])))],
    ])

    if eventos.get("audit_cleared"):
        _add_p(RC.EVENTO_1102_NOTA, bold=True)

    # --- CONCLUSIONES ---
    doc.add_page_break()
    _add_h1("7. Conclusiones")
    _add_p(RC.CONCLUSIONES_INTRO)
    for i, h in enumerate(sec["hallazgos"], 1):
        doc.add_paragraph(f"{i}. {h}", style="List Number")

    doc.add_paragraph()
    _add_p(RC.CONCLUSIONES_TRIAJE)
    doc.add_paragraph()
    _add_p(RC.CIERRE_PERICIAL.format(
        lugar_firma=sec['perito'].get('lugar_firma', '_____________'),
        fecha_emision=sec['fecha_emision'],
    ))
    doc.add_paragraph()
    doc.add_paragraph()
    _add_p(f"Fdo.: {nombre}")
    _add_p(f"{tit} — Nº Colegiado: {col}")

    # --- ANEXOS ---
    doc.add_page_break()
    _add_h1("8. Anexos")
    _add_p(RC.ANEXOS_INTRO)
    _add_p("Glosario rápido:", bold=True)
    for term, desc in RC.GLOSARIO:
        p = doc.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(desc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
